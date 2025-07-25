import os
import io
import sys
import re
import zipfile
import difflib
import shutil
import sqlite3
import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document

import copy
import pandas as pd
from docx.document import Document as _Document
from docx.table import Table
from docx.shared import Pt

#import pytz

# === 配置区域 === #
# ZIP_FILE_PATH = 'data/202504巡检记录.zip'
EXTRACT_FOLDER = 'output/unzipped_files'
DB_PATH = 'db/files_info.db'
ZIP_FILE_PATH = None  # 最终结果
DATA_DIR = 'data'
TEMPLATE_FILE = os.path.join(DATA_DIR, 'ZH-GBase8a集群-月度巡检报告-模板.docx')
TEMPLATE_FILE_OUT = Path("output/reports")
TEMPLATE_FILE_OUT.mkdir(parents=True, exist_ok=True)
# 全局共享字典
report_data = {}
#shanghai_time = datetime.now(pytz.timezone('Asia/Shanghai')).strftime('%Y-%m-%d %H:%M:%S')
# === 配置区域结束 === #

def list_zip_files():
    """列出data目录下所有zip文件"""
    if not os.path.isdir(DATA_DIR):
        return []
    return [f for f in os.listdir(DATA_DIR) if f.lower().endswith('.zip')]

def choose_from_multiple(zip_files):
    """让用户从多个zip中选择"""
    print("\n检测到多个ZIP文件，请选择要使用的文件：")
    for idx, name in enumerate(zip_files, 1):
        print(f"{idx}. {name}")
    while True:
        choice = input("请输入编号选择文件：").strip()
        if choice.isdigit() and 1 <= int(choice) <= len(zip_files):
            return os.path.join(DATA_DIR, zip_files[int(choice) - 1])
        print("无效选择，请重新输入编号。")

def manual_input_zip():
    while True:
        path = input("请输入一个有效的ZIP文件路径：").strip()
        # 去除路径前后多余引号（单引号或双引号）
        if (path.startswith('"') and path.endswith('"')) or (path.startswith("'") and path.endswith("'")):
            path = path[1:-1]

        # 将Windows反斜杠转成斜杠（跨平台都适用）
        path = path.replace('\\', '/')

        if os.path.isfile(path) and path.lower().endswith('.zip'):
            return path
        else:
            print("路径无效或不是ZIP文件，请重新输入。")

def confirm_if_not_expected(filename):
    """文件名不包含‘巡检记录’时，需要用户确认"""
    if '巡检记录' not in filename:
        print(f"\n警告：你选择的文件名 “{filename}” 不包含 '巡检记录' 字样。")
        confirm = input("是否继续使用该文件？请输入 yes 或 y 确认：").strip().lower()
        if confirm not in ['y', 'yes']:
            print("操作已取消，请重新运行。")
            exit(1)

def get_zip_file_path():
    """主逻辑"""
    global ZIP_FILE_PATH
    zip_files = list_zip_files()

    if len(zip_files) == 1:
        chosen = os.path.join(DATA_DIR, zip_files[0])
        print(f"检测到一个ZIP文件，使用：{chosen}")
    elif len(zip_files) > 1:
        chosen = choose_from_multiple(zip_files)
    else:
        chosen = manual_input_zip()

    confirm_if_not_expected(os.path.basename(chosen))
    ZIP_FILE_PATH = chosen
    print(f"\n最终使用的ZIP文件路径：{ZIP_FILE_PATH}")
    return ZIP_FILE_PATH

def initialize_project_directories():
    """初始化项目所需的目录结构，如果不存在则创建。"""
    paths_to_create = [
        os.path.dirname(ZIP_FILE_PATH) if ZIP_FILE_PATH else DATA_DIR,
        EXTRACT_FOLDER,
        os.path.dirname(DB_PATH)
    ]

    for path in paths_to_create:
        if path and not os.path.exists(path):
            os.makedirs(path, exist_ok=True)
            print(f"已创建目录：{path}")
        else:
            print(f"目录已存在：{path}")

def clear_folder(folder_path):
    """删除文件夹内所有内容，但保留文件夹本身"""
    if not os.path.exists(folder_path):
        return
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path) or os.path.islink(file_path):
            os.unlink(file_path)  # 删除文件或链接
        elif os.path.isdir(file_path):
            shutil.rmtree(file_path)  # 删除文件夹及其内容

def extract_zip(zip_path, extract_to):
    """解压 ZIP 文件，若目标文件夹非空先清空"""
    if not os.path.exists(zip_path):
        raise FileNotFoundError(f"找不到压缩包: {zip_path}")
    
    os.makedirs(extract_to, exist_ok=True)

    # 判断目标文件夹是否为空，不空则清空
    if os.listdir(extract_to):
        print(f"[!] 目标文件夹 {extract_to} 非空，正在清空...")
        clear_folder(extract_to)

    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        zip_ref.extractall(extract_to)
    
    print(f"[✓] 解压完成，文件解压到：{extract_to}")


def get_all_files(root_folder):
    """获取解压后的所有文件路径（相对路径）"""
    file_list = []
    for root, _, files in os.walk(root_folder):
        for file in files:
            rel_path = os.path.relpath(os.path.join(root, file), root_folder)
            file_list.append(rel_path)
    return file_list

def extract_system_name(filename):
    """从文件名中提取 system_name（__ 前面的部分）"""
    return filename.split("__")[0] if "__" in filename else ""

def init_database(db_path):
    """初始化 SQLite 数据库和表结构"""
    conn = sqlite3.connect(db_path)

    cursor = conn.cursor()
    # 删除旧表（如果存在）
    cursor.execute("DROP TABLE IF EXISTS files")
    cursor.execute("DROP TABLE IF EXISTS machines")
    cursor.execute("DROP TABLE IF EXISTS machine_using")
    cursor.execute("DROP TABLE IF EXISTS clusters_disk_using")
    cursor.execute("DROP TABLE IF EXISTS clusters_process")
    cursor.execute("DROP TABLE IF EXISTS clusters_logs")
    cursor.execute("DROP TABLE IF EXISTS auto_start")
    cursor.execute("DROP TABLE IF EXISTS cluster_variables")
    cursor.execute("DROP TABLE IF EXISTS data_clusters")
    cursor.execute("DROP TABLE IF EXISTS instances")
    cursor.execute("DROP TABLE IF EXISTS sys_clusters")

    # 文件记录表
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT,
            filename TEXT NOT NULL,
            fullpath TEXT NOT NULL,
            imported_at TEXT DEFAULT (datetime('now', 'localtime'))
        )
    ''')

    # 机器信息表
    cursor.execute('''
        CREATE TABLE machines (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT NOT NULL,
            cluster_name TEXT NOT NULL,
            ip_address TEXT NOT NULL,
            os_version TEXT,
            hostname TEXT,
            cpu_model_name TEXT,
            cpu_logic_core INTEGER,
            cpu_physical_core INTEGER,
            serverip_list TEXT,
            notes TEXT,
            UNIQUE(system_name,ip_address)
        );
    ''')

    # 机器资源使用表
    cursor.execute('''
        CREATE TABLE machine_using (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT NOT NULL,
            cluster_name TEXT NOT NULL,
            ip_address TEXT NOT NULL,
            mem_total TEXT,
            mem_used TEXT,
            mem_free TEXT,
            msm_shared TEXT,
            mem_buff_cache TEXT,
            mem_available TEXT,
            swap_total TEXT,
            swap_used TEXT,
            swap_free TEXT,
            disk_filesystem TEXT,
            disk_size TEXT,
            disk_used TEXT,
            disk_avail TEXT,
            disk_use_per TEXT,
            disk_mounted TEXT,
            notes TEXT,
            UNIQUE(system_name,ip_address)
        );
    ''')

    # 集群磁盘资源使用表
    cursor.execute('''
        CREATE TABLE clusters_disk_using (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT NOT NULL,
            cluster_name TEXT NOT NULL,
            disk_total TEXT,
            disk_used TEXT,
            disk_avail TEXT,
            disk_use_per TEXT,
            notes TEXT
        );
    ''')

    # 集群进程表
    cursor.execute('''
        CREATE TABLE clusters_process (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT NOT NULL,
            cluster_name TEXT NOT NULL,
            ip_address TEXT,
            process_cmd TEXT,
            notes TEXT
        );
    ''')

    # 集群进程表
    cursor.execute('''
        CREATE TABLE clusters_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT NOT NULL,
            cluster_name TEXT NOT NULL,
            ip_address TEXT,
            log_used TEXT,
            log_path TEXT,
            notes TEXT
        );
    ''')

    # 集群进程表
    cursor.execute('''
        CREATE TABLE auto_start (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT NOT NULL,
            cluster_name TEXT NOT NULL,
            ip_address TEXT,
            process_start TEXT,
            notes TEXT
        );
    ''')

    # 集群参数表
    cursor.execute('''
        CREATE TABLE cluster_variables (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT NOT NULL,
            cluster_name TEXT NOT NULL,
            ip_address TEXT,
            var_name TEXT,
            var_reference TEXT,
            config_file TEXT,
            var_actual TEXT,
            notes TEXT
        );
    ''')

    # 计算集群信息表
    cursor.execute('''
        CREATE TABLE data_clusters (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT NOT NULL,
            cluster_name TEXT NOT NULL,
            cluster_state TEXT,
            cluster_mode TEXT,
            databases_count TEXT,
            tables_count TEXT,
            views_count TEXT,
            procs_count TEXT,
            funcs_count TEXT,
            ddl_event TEXT,
            dml_event TEXT,
            dmlstorage_event TEXT,
            notes TEXT
        );
    ''')

    # 实例表
    cursor.execute('''
        CREATE TABLE instances (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT NOT NULL,
            cluster_name TEXT NOT NULL,
            namenode TEXT,
            ip_address TEXT,
            gcware TEXT,
            gcluster TEXT,
            gnode TEXT,
            syncserver TEXT,
            datastate TEXT,
            notes TEXT
        );
    ''')

    # 系统信息表
    cursor.execute('''
        CREATE TABLE sys_clusters (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            system_name TEXT NOT NULL,
            ma_one_ip TEXT,
            gbase_version TEXT,
            failover_info TEXT,
            crontab_always TEXT,
            notes TEXT
        );
    ''')

    conn.commit()
    conn.close()
    print("[✓] 数据库初始化完成")


def insert_files_to_db(db_path, file_list, extract_root):
    """将文件路径插入数据库"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    for file in file_list:
        full_path = os.path.join(extract_root, file)
        system_name = extract_system_name(file)
        cursor.execute("INSERT INTO files (system_name, filename, fullpath) VALUES (?, ?, ?)", (system_name, file, full_path))

    conn.commit()
    conn.close()
    print(f"[✓] 成功写入 {len(file_list)} 个文件记录到数据库")

def process_each_file_from_db(db_path):
    """从 files 表中读取每个 fullpath，并逐个处理文件"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    cursor.execute("SELECT id, filename, fullpath FROM files")
    rows = cursor.fetchall()

    for row in rows:
        file_id, filename, fullpath = row

        # 这里是你要对每个文件执行的操作
        try:
            with open(fullpath, 'r', encoding='utf-8') as f:
                print(f"[{file_id}] 读取文件 {filename} 成功：{fullpath}, ")
                # 你可以在这里做进一步分析
                operating_file(fullpath)
                


        except FileNotFoundError:
            print(f"[{file_id}] ❌ 文件不存在：{fullpath}")
        except Exception as e:
            print(f"[{file_id}] ⚠️ 处理异常：{e}")

    conn.close()

# ====== 文件内容处理函数 ====== #
def read_file(file_path):
    # 读取文件内容
    with open(file_path, 'r', encoding="utf-8") as file:
        #return file.readlines()
        return file.read()

def extract_blocks(text, pattern):
    # 使用 re.findall 查找匹配项
    match = re.search(pattern, text, re.DOTALL | re.MULTILINE)
    return match.group(1).strip() if match else None

def get_cluster_name(file_content):
    # 获取集群名
    first_line = file_content.splitlines()[0]
    specific_char = "GBase 8a Cluster"
    result = first_line.split(specific_char)[0]
    result = result.rstrip()
    return result

def extract_vc_names(text):
    # 获取VC名称
    pattern = r"=+ Data Machine Information\s+'([^']+)'"
    return re.findall(pattern, text)

def extract_ip_tuples(text, system_name):
    """提取IP地址和集群名的元组列表"""
    results = []
    lines = text.strip().splitlines()
    
    for line in lines:
        match = re.match(r"(\w+)\s+(\d{1,3}(?:\.\d{1,3}){3}):", line)
        if match:
            cluster_name = match.group(1)   # e.g., 'coor'
            ip_address = match.group(2)     # e.g., '10.174.68.202'
            results.append((system_name, cluster_name, ip_address))
    
    return results

def get_machine_line(text, pattern1, pattern2):
    coor_text = extract_blocks(text, pattern1)
    coor_2_text = extract_blocks(coor_text, pattern2)
    return coor_2_text.strip() if coor_2_text else None

def get_coor_ip(text):
    """提取协调节点的IP地址"""
    coor_pattern = r"^=+Coordinator Machine Information=+\n(.*?)(?=^=+Coordinator GBase Cluster Information=+)"
    coor_ip_pattern = r"\* 管理节点操作系统版本：\n(.*?)(?=\* )"
    return get_machine_line(text, coor_pattern, coor_ip_pattern)

def get_data_node_ip(text, vc_name):
    """提取数据节点的IP地址"""
    # 使用正则表达式匹配 VC 名称下的数据节点 IP 地址
    data_pattern = (
        rf"=+ Data Machine Information\s+'{vc_name}'\s+=+\n"  # 开始
        r"(.*?)"                                              # 捕获中间内容
        rf"=+ Data GBase Cluster Information\s+'{vc_name}'\s+=+"  # 结束
    )
    data_ip_pattern = r"\* 计算节点操作系统版本：\n(.*?)\* "
    return get_machine_line(text, data_pattern, data_ip_pattern)

def insert_ip_to_db(text, db_path):
    """提取IP地址 并插入到数据库"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    system_name = get_cluster_name(text)

    coor_ip_text = get_coor_ip(text)
    coor_ips = extract_ip_tuples(coor_ip_text, system_name)

    # 插入数据
    print(f"[✓] 正在处理 coor")
    cursor.executemany("INSERT OR IGNORE INTO machines (system_name, cluster_name, ip_address) VALUES (?, ?, ?)", coor_ips)
    cursor.executemany("INSERT OR IGNORE INTO machine_using (system_name, cluster_name, ip_address) VALUES (?, ?, ?)", coor_ips)

    """提取数据节点的IP地址并插入到数据库"""
    vc_names = extract_vc_names(text)

    for vc in vc_names:
        print(f"[✓] 正在处理 VC：{vc}")
        data_ip_text = get_data_node_ip(text, vc)
        data_ips = extract_ip_tuples(data_ip_text, system_name)
        # 插入数据
        cursor.executemany("INSERT OR IGNORE INTO machines (system_name, cluster_name, ip_address) VALUES (?, ?, ?)", data_ips)
        cursor.executemany("INSERT OR IGNORE INTO machine_using (system_name, cluster_name, ip_address) VALUES (?, ?, ?)", data_ips)

    conn.commit()
    conn.close()
    print(f"[✓] 成功写入IP记录到数据库")

# === 机器信息提取开始 === #
def extract_cluster_ip_machie_info(text):
    # 只匹配 IP 和 后面的信息
    pattern = r"\S+\s+([\d.]+):\s+(.*)"
    matches = re.findall(pattern, text)

    # 转成 list，再封装成 tuple 返回
    # result = [list(match) for match in matches]
    result = [[os, ip] for ip, os in matches]
    return tuple(result)

def merge_by_ip_multi(*lists):
    merged_dict = {}

    for lst in lists:
        for item in lst:
            ip = item[-1]
            values = item[:-1]
            if ip in merged_dict:
                merged_dict[ip].extend(values)
            else:
                merged_dict[ip] = values[:]

    # 最后再把 IP 加回末尾
    return [data + [ip] for ip, data in merged_dict.items()]


def get_coor_os(text):
    """提取协调节点的系统版本信息"""
    # 操作系统
    coor_pattern = r"^=+Coordinator Machine Information=+\n(.*?)(?=^=+Coordinator GBase Cluster Information=+)"
    coor_os_pattern = r"\* 管理节点操作系统版本：\n(.*?)(?=\* )"
    os_lines = get_machine_line(text, coor_pattern, coor_os_pattern)
    os_versions = extract_cluster_ip_machie_info(os_lines) if os_lines else None

    # 主机名
    coor_hostname_pattern = r"\* Hostname：\n(.*?)(?=\* )"
    hostname_lines = get_machine_line(text, coor_pattern, coor_hostname_pattern)
    hostname = extract_cluster_ip_machie_info(hostname_lines) if os_lines else None

    # cpu model name
    coor_cpu1_pattern = r"\* CPU model name信息:\n(.*?)(?=\* )"
    cpu1_lines = get_machine_line(text, coor_pattern, coor_cpu1_pattern)
    cpu1 = extract_cluster_ip_machie_info(cpu1_lines) if os_lines else None
    # cpu1 = extract_cluster_ip_machie_info(cpu1_lines) if os_lines else None
    
    # cpu 逻辑核心数
    coor_cpu2_pattern = r"\* CPU 逻辑核数信息：\n(.*?)(?=\* )"
    cpu2_lines = get_machine_line(text, coor_pattern, coor_cpu2_pattern)
    cpu2 = extract_cluster_ip_machie_info(cpu2_lines) if os_lines else None

    # cpu 物理核心数
    coor_cpu3_pattern = r"\* CPU 物理核数：\n(.*?)(?=\* )"
    cpu3_lines = get_machine_line(text, coor_pattern, coor_cpu3_pattern)
    cpu3 = extract_cluster_ip_machie_info(cpu3_lines) if os_lines else None 

    # serverip_list
    """目前有问题，该列置空，正常再添加"""
    coor_iplist_pattern = r"\* 服务器IP地址列表\n(.*?)(?=\* )"
    # iplist_lines = get_machine_line(text, coor_pattern, coor_iplist_pattern)
    # iplist = extract_cluster_ip_machie_info(iplist_lines) if os_lines else None 

    # 下面是置空模拟，正常后需要修改
    pattern_iplist_name = r"coor\s+(\d{1,3}(?:\.\d{1,3}){3}):"
    iplist_lines = get_machine_line(text, coor_pattern, coor_iplist_pattern)
    iplist_model = re.findall(pattern_iplist_name, iplist_lines)
    unique_ips = sorted(set(iplist_model)) # 去重并排序
    iplist = [['', ip] for ip in unique_ips]

    # 合并所有信息
    merged = merge_by_ip_multi(os_versions, hostname, cpu1, cpu2, cpu3, iplist)
    return merged

def get_data_node_info(text, vc_name):
    """提取数据节点的系统版本信息"""
    
    data_pattern = (
        rf"=+ Data Machine Information\s+'{vc_name}'\s+=+\n"  # 开始
        r"(.*?)"                                              # 捕获中间内容
        rf"=+ Data GBase Cluster Information\s+'{vc_name}'\s+=+"  # 结束
    )
    data_os_pattern = r"\* 计算节点操作系统版本：\n(.*?)\* "
    os_lines = get_machine_line(text, data_pattern, data_os_pattern)
    os_versions = extract_cluster_ip_machie_info(os_lines) if os_lines else None    

    # 主机名
    data_hostname_pattern = r"\* Hostname：\n(.*?)(?=\* )"
    hostname_lines = get_machine_line(text, data_pattern, data_hostname_pattern)
    hostname = extract_cluster_ip_machie_info(hostname_lines) if os_lines else None

    # cpu model name
    data_cpu1_pattern = r"\* CPU model name信息:\n(.*?)(?=\* )"
    cpu1_lines = get_machine_line(text, data_pattern, data_cpu1_pattern)
    cpu1 = extract_cluster_ip_machie_info(cpu1_lines) if os_lines else None
    # cpu1 = extract_cluster_ip_machie_info(cpu1_lines) if os_lines else None
    
    # cpu 逻辑核心数
    data_cpu2_pattern = r"\* CPU 逻辑核数：\n(.*?)(?=\* )"
    cpu2_lines = get_machine_line(text, data_pattern, data_cpu2_pattern)
    cpu2 = extract_cluster_ip_machie_info(cpu2_lines) if os_lines else None


    # cpu 物理核心数
    data_cpu3_pattern = r"\* CPU 物理核数：\n(.*?)(?=\* )"
    cpu3_lines = get_machine_line(text, data_pattern, data_cpu3_pattern)
    cpu3 = extract_cluster_ip_machie_info(cpu3_lines) if os_lines else None 

    # serverip_list
    """目前有问题，该列置空，正常再添加"""
    data_iplist_pattern = r"\* 计算集群IP列表：\n(.*?)(?=\* )"
    # iplist_lines = get_machine_line(text, data_pattern, data_iplist_pattern)
    # iplist = extract_cluster_ip_machie_info(iplist_lines) if os_lines else None 

    # 下面是置空模拟，正常后需要修改
    pattern_iplist_name = r"\s+(\d{1,3}(?:\.\d{1,3}){3}):"
    iplist_lines = get_machine_line(text, data_pattern, data_iplist_pattern)
    iplist_model = re.findall(pattern_iplist_name, iplist_lines)
    unique_ips = sorted(set(iplist_model)) # 去重并排序
    iplist = [['', ip] for ip in unique_ips]
    # print(f"获取到 {vc_name} 的IP地址列表：{iplist}")

    # 合并所有信息
    merged = merge_by_ip_multi(os_versions, hostname, cpu1, cpu2, cpu3, iplist)
    return merged


def get_machine_info(text, db_path):
    """从文件内容中提取机器信息"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 提取协调节点信息
    get_coor_info = get_coor_os(text)

    # coor节点的os信息插入数据库
    print(f"[✓] 正在处理 coor：")
    cursor.executemany(
        """
        UPDATE machines 
        SET os_version = ?, hostname = ?, cpu_model_name = ?, cpu_logic_core = ?, cpu_physical_core = ?, serverip_list = ?
        WHERE ip_address = ?
        """, 
        get_coor_info)
    
    """提取数据节点的IP地址并插入到数据库"""
    vc_names = extract_vc_names(text)

    for vc in vc_names:
        print(f"[✓] 正在处理 VC：{vc}")
        data_info = get_data_node_info(text, vc)
        # print(f"获取到 {vc} 的数据节点信息：{data_info}")
        # 插入数据
        cursor.executemany(
            """
            UPDATE machines 
            SET os_version = ?, hostname = ?, cpu_model_name = ?, cpu_logic_core = ?, cpu_physical_core = ?, serverip_list = ?
            WHERE ip_address = ?
            """, 
            data_info)

    conn.commit()
    conn.close()
    print(f"[✓] 成功写入数据节点的机器信息到数据库")

# === machine信息提取结束 === #

# === machine memory,swap,disk 信息提取开始 === #
def extract_mem_info(text):
    mem_list = []
    lines = text.strip().splitlines()

    for line in lines:
        match = re.match(r".*?(\d+\.\d+\.\d+\.\d+):\s*Mem:\s*(.+)", line)
        if match:
            ip = match.group(1)
            numbers = match.group(2).split()
            numbers.append(ip)
            mem_list.append(numbers)
    return mem_list

def extract_swap_info(text):
    swap_list = []
    lines = text.strip().splitlines()

    for line in lines:
        match = re.match(r".*?(\d+\.\d+\.\d+\.\d+):\s*Swap:\s*(.+)", line)
        if match:
            ip = match.group(1)
            numbers = match.group(2).split()
            numbers.append(ip)
            swap_list.append(numbers)
    return swap_list

def extract_df_info(text):
    df_list = []
    lines = text.strip().splitlines()

    for line in lines:
        match = re.match(r".*?(\d+\.\d+\.\d+\.\d+):\s*(.+)", line)
        if match:
            ip = match.group(1)
            df_data = match.group(2).split()
            if len(df_data) >= 6:
                df_data.append(ip)
                df_list.append(df_data)
    return df_list

def get_coor_using(text):
    """提取协调节点的内存，swap，磁盘使用信息"""

    coor_pattern = r"^=+Coordinator Machine Information=+\n(.*?)(?=^=+Coordinator GBase Cluster Information=+)"
    coor_mem_pattern = r"\ * 物理内存使用情况：\n(.*?)(?=\* )"
    coor_swap_pattern = r"\ * SWAP内存使用情况：\n(.*?)(?=\* )"
    coor_disk_pattern = r"\ * 管理节点空间使用情况：\n(.*?)(?=\* )"

    mem_lines = get_machine_line(text, coor_pattern, coor_mem_pattern)
    swap_lines = get_machine_line(text, coor_pattern, coor_swap_pattern)
    disk_lines = get_machine_line(text, coor_pattern, coor_disk_pattern)

    mem_use_info = extract_mem_info(mem_lines) if mem_lines else None
    swap_use_info = extract_swap_info(swap_lines) if swap_lines else None
    disk_use_info = extract_df_info(disk_lines) if disk_lines else None

    # 合并所有信息
    merged = merge_by_ip_multi(mem_use_info, swap_use_info, disk_use_info)
    return merged

def get_data_node_using(text, vc_name):
    """提取数据节点的内存，swap，磁盘使用信息"""
    # coor_pattern = r"^=+Coordinator Machine Information=+\n(.*?)(?=^=+Coordinator GBase Cluster Information=+)"
    data_pattern = (
        rf"=+ Data Machine Information\s+'{vc_name}'\s+=+\n"  # 开始
        r"(.*?)"                                              # 捕获中间内容
        rf"=+ Data GBase Cluster Information\s+'{vc_name}'\s+=+"  # 结束
    )
    mem_pattern = r"\ * 物理内存使用情况\n(.*?)(?=\* )"
    swap_pattern = r"\ * SWAP使用情况：\n(.*?)(?=\* )"
    disk_pattern = r"\ * 计算集群各节点空间情况：\n(.*?)(?=\* )"

    mem_lines = get_machine_line(text, data_pattern, mem_pattern)
    swap_lines = get_machine_line(text, data_pattern, swap_pattern)
    disk_lines = get_machine_line(text, data_pattern, disk_pattern)

    mem_use_info = extract_mem_info(mem_lines) if mem_lines else None
    swap_use_info = extract_swap_info(swap_lines) if swap_lines else None
    disk_use_info = extract_df_info(disk_lines) if disk_lines else None

    # 合并所有信息
    merged = merge_by_ip_multi(mem_use_info, swap_use_info, disk_use_info)
    return merged

def get_machine_using(text, db_path):
    """从文件内容中提取机器使用信息"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 提取协调节点信息
    get_coor_info = get_coor_using(text)

    # coor节点的os信息插入数据库
    print(f"[✓] 正在处理 coor：")
    cursor.executemany(
        """
        UPDATE machine_using 
        SET
        mem_total = ? ,
        mem_used = ? ,
        mem_free = ? ,
        msm_shared = ? ,
        mem_buff_cache = ? ,
        mem_available = ? ,
        swap_total = ? ,
        swap_used = ? ,
        swap_free = ? ,
        disk_filesystem = ? ,
        disk_size = ? ,
        disk_used = ? ,
        disk_avail = ? ,
        disk_use_per = ? ,
        disk_mounted = ?
        WHERE ip_address = ?
        """, 
        get_coor_info)
    
    """提取数据节点的IP地址并插入到数据库"""
    vc_names = extract_vc_names(text)

    for vc in vc_names:
        print(f"[✓] 正在处理 VC：{vc}")
        data_info = get_data_node_using(text, vc)
        # 插入数据
        cursor.executemany(
            """
            UPDATE machine_using 
            SET
            mem_total = ? ,
            mem_used = ? ,
            mem_free = ? ,
            msm_shared = ? ,
            mem_buff_cache = ? ,
            mem_available = ? ,
            swap_total = ? ,
            swap_used = ? ,
            swap_free = ? ,
            disk_filesystem = ? ,
            disk_size = ? ,
            disk_used = ? ,
            disk_avail = ? ,
            disk_use_per = ? ,
            disk_mounted = ?
            WHERE ip_address = ?
            """, 
            data_info)

    conn.commit()
    conn.close()
    print(f"[✓] 成功写入数据节点的机器内存，swap，disk使用信息到数据库")

# === machine memory,swap,disk 信息提取结束 === #

# === 集群磁盘使用信息提取开始 === #
def extract_next_line_values_list(text: str, keywords: list, prefix: list) -> list:
    lines = text.strip().splitlines()
    results = []
    for keyword in keywords:
        for i, line in enumerate(lines):
            if keyword in line:
                if i + 1 < len(lines):
                    results.append(lines[i + 1].strip())
                else:
                    results.append('')  # 没有下一行，填空
                break
    return prefix + results



def get_cluster_disk_using(text, db_path):
    """提取集群磁盘使用信息"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 提取协调节点信息
    coor_pattern = r"^=+Coordinator Machine Information=+\n(.*?)(?=^=+Coordinator GBase Cluster Information=+)"
    disk_text = extract_blocks(text, coor_pattern)
    keywords = ['管理节点总空间之和', '管理节点已使用空间之和', '管理节点剩余用空间之和', '管理集群空间总使用率']

    system_name = get_cluster_name(text)
    prefix = [system_name, 'coor']
    disk_result = extract_next_line_values_list(disk_text, keywords, prefix)
    # values = extract_next_line_values_list(disk_text, keywords)
    cursor.executemany("INSERT OR IGNORE INTO clusters_disk_using (system_name, cluster_name, disk_total, disk_used, disk_avail, disk_use_per) VALUES (?, ?, ?, ?, ?, ?)", [disk_result])
    
    """提取数据节点的磁盘使用信息"""
    vc_names = extract_vc_names(text)

    for vc in vc_names:
        print(f"[✓] 正在处理 VC：{vc}")
        data_pattern = (
            rf"=+ Data Machine Information\s+'{vc}'\s+=+\n"  # 开始
            r"(.*?)"                                              # 捕获中间内容
            rf"=+ Data GBase Cluster Information\s+'{vc}'\s+=+"  # 结束
        )
        disk_text = extract_blocks(text, data_pattern)
        keywords = ['计算集群空间之和', '计算集群已使用空间之和', '计算集群剩余用空间之和', '计算集群空间总使用率']
        prefix = [system_name, vc]
        disk_result2 = extract_next_line_values_list(disk_text, keywords, prefix)
        # print(f"[✓] 正在处理计算集群磁盘使用信息：{disk_result2}")
        cursor.executemany("INSERT OR IGNORE INTO clusters_disk_using (system_name, cluster_name, disk_total, disk_used, disk_avail, disk_use_per) VALUES (?, ?, ?, ?, ?, ?)", [disk_result2])

    conn.commit()
    conn.close()
    print(f"[✓] 成功写入集群磁盘使用信息到数据库")

# === 集群磁盘使用信息提取结束 === #


# === 提取节点进程 === #
def extract_ip_process_pairs(text):
    blocks = re.split(r'-{5,}\s*(\d+\.\d+\.\d+\.\d+)\s*-{5,}', text)
    results = []

    for i in range(1, len(blocks), 2):
        ip = blocks[i]
        content = blocks[i + 1]

        # 匹配 ps -ef 输出中每一行的最后一个字段（即命令部分）
        lines = content.strip().splitlines()
        for line in lines:
            # 避免空行或异常行
            if not line.strip():
                continue

            # 以空格拆分行，前 7 项是 ps 输出的标准字段（uid, pid, ppid, etc），之后都是命令
            parts = line.split(None, 7)  # 最多分成 8 段
            if len(parts) < 8:
                continue  # 不足 8 段说明不是标准 ps -ef 行

            command = parts[7].strip()
            results.append([ip, command])

    return results


def extract_ip_command_pairs_linewise(text):
    result = []
    lines = text.strip().splitlines()

    for line in lines:
        ip_match = re.search(r'(\d+\.\d+\.\d+\.\d+):', line)
        if not ip_match:
            continue
        ip = ip_match.group(1)

        # 获取IP后面的内容
        after_ip = line.split(f'{ip}:', 1)[1].strip()

        # ps -ef字段一般固定，取前7字段，剩余是命令
        fields = after_ip.split(None, 7)
        if len(fields) < 8:
            # 没有足够字段，跳过
            continue

        command = fields[7].strip()
        result.append([ip, command])

    return result

def get_cluster_process(text, db_path):
    """提取集群节点进程信息"""
    # 这里可以添加对进程信息的提取逻辑
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 提取协调节点信息
    coor_pattern = r"^=+Coordinator GBase Cluster Information=+\n(.*?)(?=^=+Coordinator GBase Cluster variables=+)"
    process_pattern = r"\ * 管理节点进程状态：\n(.*?)(?=\* 管理节点日志大小)"
    coor_text = extract_blocks(text, coor_pattern)
    process_line = extract_blocks(coor_text, process_pattern)
    process_list = extract_ip_process_pairs(process_line)

    system_name = get_cluster_name(text)
    prefix = [system_name, 'coor']
    process_list_prefix = [prefix + row for row in process_list]
    tuple_data = [tuple(row) for row in process_list_prefix]
  

    # print(f"[✓] 正在处理集群磁盘使用信息：{tuple_data}")
    cursor.executemany("INSERT OR IGNORE INTO clusters_process (system_name, cluster_name, ip_address, process_cmd) VALUES (?, ?, ?, ?)", tuple_data)
    
    vc_names = extract_vc_names(text)

    for vc in vc_names:
        print(f"[✓] 正在处理 VC：{vc}")
        data_pattern = (
            rf"=+ Data GBase Cluster Information\s+'{vc}'\s+=+\n"  # 开始
            r"(.*?)"                                              # 捕获中间内容
            rf"=+ Data GBase Cluster variables\s+{vc}\s+=+"  # 结束
        )
        data_process_pattern = r"\* Data Cluster 进程状态:\s*(.*?)\s*\* Data Cluster 日志情况:"
        data_text = extract_blocks(text, data_pattern)
        process_line2 = extract_blocks(data_text, data_process_pattern)
        process_list2 = extract_ip_command_pairs_linewise(process_line2)
        prefix2 = [system_name, vc]
        process_list_prefix2 = [prefix2 + row for row in process_list2]
        tuple_data2 = [tuple(row) for row in process_list_prefix2]
        cursor.executemany("INSERT OR IGNORE INTO clusters_process (system_name, cluster_name, ip_address, process_cmd) VALUES (?, ?, ?, ?)", tuple_data2)

    conn.commit()
    conn.close()
    print(f"[✓] 成功写入集群进程信息到数据库")

# === 提取节点进程结束 === #

# === 提取集群日志信息 === #
def extract_du_info(text):
    result = []
    lines = text.strip().splitlines()

    for line in lines:
        # 跳过包含错误信息的行
        if "No such file or directory" in line:
            continue

        ip_match = re.search(r'(\d+\.\d+\.\d+\.\d+):', line)
        if not ip_match:
            continue
        ip = ip_match.group(1)

        # 提取大小和路径
        match = re.search(r':\s+(\S+)\s+(.+)', line)
        if match:
            size = match.group(1)
            path = match.group(2).strip()
            result.append([ip, size, path])

    return result

def get_cluster_logs(text, db_path):
    """提取集群节点日志信息"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 提取协调节点信息
    coor_pattern = r"^=+Coordinator GBase Cluster Information=+\n(.*?)(?=^=+Coordinator GBase Cluster variables=+)"
    process_pattern = r"\ * 管理节点日志大小：\n(.*?)(?=\ * 自启动设置：)"
    coor_text = extract_blocks(text, coor_pattern)
    logs_line = extract_blocks(coor_text, process_pattern)
    logs_list = extract_du_info(logs_line)

    system_name = get_cluster_name(text)
    prefix = [system_name, 'coor']
    logs_list_prefix = [prefix + row for row in logs_list]
    tuple_data = [tuple(row) for row in logs_list_prefix]
  
    cursor.executemany("INSERT OR IGNORE INTO clusters_logs (system_name, cluster_name, ip_address, log_used, log_path) VALUES (?, ?, ?, ?, ?)", tuple_data)
    
    """提取数据节点的磁盘使用信息"""
    vc_names = extract_vc_names(text)

    for vc in vc_names:
        print(f"[✓] 正在处理 VC：{vc}")
        data_pattern = (
            rf"=+ Data GBase Cluster Information\s+'{vc}'\s+=+\n"  # 开始
            r"(.*?)"                                              # 捕获中间内容
            rf"=+ Data GBase Cluster variables\s+{vc}\s+=+"  # 结束
        )
        data_logs_pattern = r"\* Data Cluster 日志情况:\s*(.*?)\s*\* Data Cluster 自启动："
        data_text = extract_blocks(text, data_pattern)
        logs_line2 = extract_blocks(data_text, data_logs_pattern)
        logs_list2 = extract_du_info(logs_line2)
        prefix2 = [system_name, vc]
        logs_list_prefix2 = [prefix2 + row for row in logs_list2]
        tuple_data2 = [tuple(row) for row in logs_list_prefix2]
        cursor.executemany("INSERT OR IGNORE INTO clusters_logs (system_name, cluster_name, ip_address, log_used, log_path) VALUES (?, ?, ?, ?, ?)", tuple_data2)
    
    conn.commit()
    conn.close()
    print(f"[✓] 成功写入集群日志信息到数据库")

# === 提取集群日志信息结束 === #

# === 提取自启动信息 === #
def extract_ip_command_pairs(text):
    result = []
    current_ip = None

    for line in text.strip().splitlines():
        line = line.strip()

        # 匹配 IP 段头：--------- <IP>---------
        ip_match = re.match(r'-+\s*(\d+\.\d+\.\d+\.\d+)\s*-+', line)
        if ip_match:
            current_ip = ip_match.group(1)
            continue

        # 非空行且已知当前 IP，就认为是命令行
        if current_ip and line:
            result.append([current_ip, line])

    return result

def extract_ip_command_data_pairs(text):

    result = []
    for line in text.strip().splitlines():
        match = re.match(r'.*?(\d+\.\d+\.\d+\.\d+):\s+(.*)', line)
        if match:
            ip = match.group(1)
            command = match.group(2)
            result.append([ip, command])
    return result

def get_auto_start(text, db_path):
    # 这里可以添加对进程信息的提取逻辑
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 提取协调节点信息
    coor_pattern = r"^=+Coordinator GBase Cluster Information=+\n(.*?)(?=^=+Coordinator GBase Cluster variables=+)"
    auto_pattern = r"\ * 自启动设置：\n(.*?)(?=\* 监控运维脚本：)"
    coor_text = extract_blocks(text, coor_pattern)
    auto_line = extract_blocks(coor_text, auto_pattern)
    # print(f"[✓] 正在处理集群自启动信息：{auto_line}")
    process_list = extract_ip_command_pairs(auto_line)
    # print(f"[✓] 处理后的自启动信息：{process_list}")

    system_name = get_cluster_name(text)
    prefix = [system_name, 'coor']
    process_list_prefix = [prefix + row for row in process_list]
    tuple_data = [tuple(row) for row in process_list_prefix]
  

    # print(f"[✓] 正在处理集群磁盘使用信息：{tuple_data}")
    cursor.executemany("INSERT OR IGNORE INTO auto_start (system_name, cluster_name, ip_address, process_start) VALUES (?, ?, ?, ?)", tuple_data)
    
    vc_names = extract_vc_names(text)

    for vc in vc_names:
        print(f"[✓] 正在处理 VC：{vc}")
        data_pattern = (
            rf"=+ Data GBase Cluster Information\s+'{vc}'\s+=+\n"  # 开始
            r"(.*?)"                                              # 捕获中间内容
            rf"=+ Data GBase Cluster variables\s+{vc}\s+=+"  # 结束
        )
        data_text = extract_blocks(text, data_pattern)
        auto_line2 = re.search(r'\* Data Cluster 自启动：\n(.*)', data_text, re.DOTALL)
        if auto_line2:
            result = auto_line2.group(1).strip()
            auto_list2 = extract_ip_command_data_pairs(result)
            prefix2 = [system_name, vc]
            process_list_prefix2 = [prefix2 + row for row in auto_list2]
            tuple_data2 = [tuple(row) for row in process_list_prefix2]
            # print(f"[✓] 处理后的自启动信息：{tuple_data2}")
            cursor.executemany("INSERT OR IGNORE INTO auto_start (system_name, cluster_name, ip_address, process_start) VALUES (?, ?, ?, ?)", tuple_data2)

    conn.commit()
    conn.close()
    print(f"[✓] 成功写入集群自启动信息到数据库")

# === 提取自启动信息结束 === #

# === 提取集群变量信息 === #
def extract_coordinator_variables_block(text):
    """
    提取协调节点的集群变量信息块，直接返回整段文本，不拆行、不做多余处理。
    """
    match = re.search(
        r'={10,}Coordinator GBase Cluster variables={10,}(.*?)GBase 8a Cluster Coordinator Inscpection End now',
        text,
        re.DOTALL
    )

    if match:
        return match.group(1).strip()  # 返回整个文本块（去首尾空白）
    else:
        return ''  # 没有匹配时返回空字符串

def extract_ip_ref_actual_params1(text):
    result = []
    current_param = None  # (参数名, 参考值)

    for line in text.strip().splitlines():
        line = line.strip()

        # 提取参考参数名和值
        ref_match = re.match(r'>>\s*([\w_]+)=([\w\.]+):', line)
        if ref_match:
            param_name = ref_match.group(1)
            ref_value = ref_match.group(2)
            current_param = (param_name, ref_value)
            continue

        # 跳过注释行
        if '#' in line:
            continue

        # 匹配格式一：带路径的配置参数
        match_with_path = re.match(
            r'coor\s+(\d+\.\d+\.\d+\.\d+):\s*(\/[^\s:]+):([\w_]+)=([\w\.]+)', line)
        if match_with_path and current_param:
            ip = match_with_path.group(1)
            filepath = match_with_path.group(2)
            param_name_actual = match_with_path.group(3)
            actual_value = match_with_path.group(4)

            if param_name_actual == current_param[0]:
                result.append([ip, param_name_actual, current_param[1], filepath, actual_value])
            continue

        # 匹配格式二：不带路径的
        match_simple = re.match(
            r'coor\s+(\d+\.\d+\.\d+\.\d+):\s*([\w_]+)=([\w\.]+)', line)
        if match_simple and current_param:
            ip = match_simple.group(1)
            param_name_actual = match_simple.group(2)
            actual_value = match_simple.group(3)

            if param_name_actual == current_param[0]:
                result.append([ip, param_name_actual, current_param[1], '', actual_value])

    return result

def extract_ip_ref_actual_params(text):
    result = []
    current_ref_param = None
    current_ref_value = None

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        # 提取参考值行，例如 >> param=value:
        ref_match = re.match(r">>\s*([a-zA-Z0-9_]+)=([^\s:]+):", line)
        if ref_match:
            current_ref_param = ref_match.group(1)
            current_ref_value = ref_match.group(2)
            continue

        # 匹配包含配置路径的行或不包含配置路径的行
        actual_match = re.match(
            r"^([a-zA-Z0-9_]+)\s+(\d+\.\d+\.\d+\.\d+):(?:\s*([^:]*):)?\s*#?([a-zA-Z0-9_]+)=([^\s]+)", line
        )

        if actual_match and current_ref_param:
            # 解析字段
            prefix = actual_match.group(1)                        # 节点前缀
            ip = actual_match.group(2)                            # IP地址
            config_path = actual_match.group(3) or ''             # 配置路径（可选）
            param_key = actual_match.group(4)                     # 参数名
            actual_value = actual_match.group(5)                  # 实际值

            # 忽略被注释的行（#param=...）
            if line.split(":")[-1].lstrip().startswith("#"):
                continue

            result.append([
                ip,
                param_key,
                current_ref_value,
                config_path.strip(),
                actual_value
            ])

    return result

def extract_data_cluster_variables_block(text, vc_name):
    data_pattern = rf"================================ Data GBase Cluster variables\s+{re.escape(vc_name)}\s+={{3,}}(.*?)GBase 8a Cluster Data Cluster '{re.escape(vc_name)}' Inscpection End"

    match = re.search(data_pattern, text, re.DOTALL)
    if match:
        return match.group(1).strip()
    else:
        return ''

def get_cluster_variables(text, db_path):
    """提取集群变量信息"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 提取协调节点信息
    system_name = get_cluster_name(text)
    coor_var_result = extract_coordinator_variables_block(text)
    coor_var_lines = extract_ip_ref_actual_params(coor_var_result)
    # print(f"[✓] 正在处理协调节点的集群变量信息：{coor_var_lines}")
    if coor_var_lines:  # 如果 data 非空
        
        prefix = [system_name, 'coor']
        var_list_prefix = [prefix + row for row in coor_var_lines]
        tuple_data = [tuple(row) for row in var_list_prefix]
        cursor.executemany("INSERT OR IGNORE INTO cluster_variables (system_name, cluster_name, ip_address, var_name, var_reference, config_file, var_actual) VALUES (?, ?, ?, ?, ?, ?, ?)", tuple_data)
    
    vc_names = extract_vc_names(text)

    for vc in vc_names:
        print(f"[✓] 正在处理 VC：{vc}")
        data_text = extract_data_cluster_variables_block(text, vc)
        data_var_lines = extract_ip_ref_actual_params(data_text)
        # print(f"[✓] 提取到 {vc} 的集群变量信息：{data_var_lines}")
        if data_var_lines:  # 如果 data 非空
        
            prefix2 = [system_name, vc]
            var_list_prefix2 = [prefix2 + row for row in data_var_lines]
            tuple_data2 = [tuple(row) for row in var_list_prefix2]
            # print(f"[✓] 正在处理 {vc} 的集群变量信息：{tuple_data2}")
            cursor.executemany("INSERT OR IGNORE INTO cluster_variables (system_name, cluster_name, ip_address, var_name, var_reference, config_file, var_actual) VALUES (?, ?, ?, ?, ?, ?, ?)", tuple_data2)
        
    conn.commit()
    conn.close()
    print(f"[✓] 成功写入集群参数变量到数据库")  
# === 提取集群变量信息结束 === #

# === 提取数据集群使用情况 === #
def extract_cluster_status(text):
    result = []

    patterns = [
        r"CLUSTER STATE:\s*([^\s\r\n]+)",             # 提取 CLUSTER STATE
        r"VIRTUAL CLUSTER MODE:\s*([^\s\r\n]+)"       # 提取 VIRTUAL CLUSTER MODE
    ]

    for pattern in patterns:
        match = re.search(pattern, text)
        result.append(match.group(1) if match else '')

    return result

def extract_vc_event_counts(text):
    # 匹配目标段落后下一行
    pattern = r"Data Cluster DDL&DML&DMLSTORAGE Event信息：\s*\n([^\n]*)"
    match = re.search(pattern, text)
    if match:
        line = match.group(1)
        # 提取该行中所有 Vc event count:数字
        return re.findall(r"Vc event count:(\d+)", line)
    else:
        return []

def get_data_cluster_using(text, db_path):
    """提取集群使用信息"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    system_name = get_cluster_name(text)
    vc_names = extract_vc_names(text)

    for vc in vc_names:
        print(f"[✓] 正在处理 VC：{vc}")
        data_pattern = (
            rf"=+ Data GBase Cluster Information\s+'{vc}'\s+=+\n"  # 开始
            r"(.*?)"                                              # 捕获中间内容
            rf"=+ Data GBase Cluster variables\s+{vc}\s+=+"  # 结束
        )
        data_text = extract_blocks(text, data_pattern)
        cluster_state = extract_cluster_status(data_text)
        event_list = extract_vc_event_counts(data_text)
        keywords = ['库的个数', '表的个数', '视图的个数', '存储过程的个数', '函数的个数']
        prefix = [system_name, vc]
        prefix2= prefix + cluster_state
        data_result2 = extract_next_line_values_list(data_text, keywords, prefix2)
        data_result3 = data_result2 + event_list
        # print(f"[✓] 正在处理 {vc} 的数据集群使用信息：{data_result3}")
        # 提取集群状态
    
        cursor.executemany("""
                           INSERT OR IGNORE INTO data_clusters (
                           system_name, cluster_name, cluster_state, cluster_mode, 
                           databases_count, tables_count, views_count, procs_count, funcs_count,
                           ddl_event, dml_event, dmlstorage_event) 
                           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                           """, [data_result3])

    conn.commit()
    conn.close()
    print(f"[✓] 成功写入集群信息到数据库")

# === 集群使用信息提取结束 === #
def clean_ansi_escape(s):
    ansi_escape = re.compile(r'\x1b\[[0-9;]*m')
    return ansi_escape.sub('', s)

def extract_gcware_and_coordinator(text):
    gcware_list = []
    coordinator_list = []

    inside_gcware = False
    inside_coordinator = False

    for line in text.splitlines():
        line = line.strip()
        line = clean_ansi_escape(line)  # 先清理整行的ANSI码

        if "GBASE GCWARE CLUSTER INFORMATION" in line:
            inside_gcware = True
            inside_coordinator = False
            continue
        elif "GBASE COORDINATOR CLUSTER INFORMATION" in line:
            inside_gcware = False
            inside_coordinator = True
            continue
        elif "GBASE VIRTUAL CLUSTER INFORMATION" in line:
            inside_gcware = False
            inside_coordinator = False
            continue

        if line.startswith("|") and line.endswith("|") and not set(line) <= set("|-="):
            parts = [p.strip() for p in line.strip('|').split('|')]

            if inside_gcware and len(parts) >= 3 and parts[0].startswith("gcware"):
                gcware_list.append(parts[:3])
            elif inside_coordinator and len(parts) >= 4 and parts[0].startswith("coordinator"):
                coordinator_list.append(parts[:4])

    return gcware_list, coordinator_list

def extract_ins_node_info(text):
    lines = clean_ansi_escape(text).splitlines()
    node_list = []
    in_table = False
    header_indexes = {}

    for i, line in enumerate(lines):
        line = line.strip()

        # 寻找表头行
        if re.match(r'\| *NodeName *\|.*\| *DataState *\|', line):
            headers = [h.strip() for h in line.strip('|').split('|')]
            for idx, header in enumerate(headers):
                header_indexes[header.lower()] = idx
            in_table = True
            continue

        if in_table:
            # 表格结束条件：不再以“|”开头
            if not line.startswith('|') or re.match(r'^[-=]+$', line.replace('|', '').strip()):
                continue

            parts = [p.strip() for p in line.strip('|').split('|')]

            try:
                name = parts[header_indexes['nodename']]
                ip = parts[header_indexes['ipaddress']]
                gnode = parts[header_indexes['gnode']]
                syncserver = parts[header_indexes['syncserver']]
                datastate = parts[header_indexes['datastate']]
                if name.startswith("node"):
                    node_list.append([name, ip, gnode, syncserver, datastate])
            except (KeyError, IndexError):
                continue  # 跳过不完整或格式不对的行

    return node_list

def get_instances(text, db_path):
    # 这里可以添加对进程信息的提取逻辑
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 提取协调节点信息
    coor_pattern = r"^=+Coordinator GBase Cluster Information=+\n(.*?)(?=^=+Coordinator GBase Cluster variables=+)"
    ins_pattern = r"\ * Coor Cluster拓扑及状态：\n(.*?)(?=\ * Coor Cluster Failover信息：)"
    coor_text = extract_blocks(text, coor_pattern)
    ins_line = extract_blocks(coor_text, ins_pattern)
    # print(f"[✓] 正在处理集群自启动信息：{ins_line}")
    gcware_info, coordinator_info = extract_gcware_and_coordinator(ins_line)
    system_name = get_cluster_name(text)
    prefix = [system_name, 'coor']
    gcware_info_prefix = [prefix + row for row in gcware_info]
    coordinator_info_prefix = [prefix + row for row in coordinator_info]

    tuple_gcware = [tuple(row) for row in gcware_info_prefix]
    tuple_coordinator = [tuple(row) for row in coordinator_info_prefix]

    cursor.executemany("INSERT OR IGNORE INTO instances (system_name, cluster_name, namenode, ip_address, gcware) VALUES (?, ?, ?, ?, ?)", tuple_gcware)
    cursor.executemany("INSERT OR IGNORE INTO instances (system_name, cluster_name, namenode, ip_address, gcluster, datastate) VALUES (?, ?, ?, ?, ?, ?)", tuple_coordinator)

    vc_names = extract_vc_names(text)

    for vc in vc_names:
        print(f"[✓] 正在处理 VC：{vc}")
        data_pattern = (
            rf"=+ Data GBase Cluster Information\s+'{vc}'\s+=+\n"  # 开始
            r"(.*?)"                                              # 捕获中间内容
            rf"=+ Data GBase Cluster variables\s+{vc}\s+=+"  # 结束
        )
        data_text = extract_blocks(text, data_pattern)
        ins_pattern2 = r"\ * Data Cluster 拓扑及状态：\n(.*?)(?=\ * Data Cluster DDL&DML&DMLSTORAGE Event信息：)"
        ins_line2 = extract_blocks(data_text, ins_pattern2)
        ins2 = extract_ins_node_info(ins_line2)

        prefix2 = [system_name, vc]
        ins_list_prefix2 = [prefix2 + row for row in ins2]
        tuple_data2 = [tuple(row) for row in ins_list_prefix2]
        cursor.executemany("""
                           INSERT OR IGNORE INTO instances (
                           system_name, cluster_name, namenode, ip_address, gnode, syncserver, datastate
                           ) 
                           VALUES (?, ?, ?, ?, ?, ?, ?)
                           """, tuple_data2)

    conn.commit()
    conn.close()
    print(f"[✓] 成功写入集群实例信息到数据库")


def extract_coordinator1_ip(text):
    clean_text = clean_ansi_escape(text)
    match = re.search(r'\| *coordinator1 *\| *([\d.]+) *\|', clean_text)
    if match:
        return match.group(1)
    return None

def extract_gbase_version(text):
    match = re.search(r'GBase版本号.*?\n([^\n]+)', text)
    if match:
        return match.group(1).strip()
    return None

def get_sys_cluster(text, db_path):
    # 这里可以添加对进程信息的提取逻辑
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # 提取协调节点信息
    coor_pattern = r"^=+Coordinator GBase Cluster Information=+\n(.*?)(?=^=+Coordinator GBase Cluster variables=+)"
    cron_pattern = r"\ * 监控运维脚本：\n(.*?)(?=\ * Coor Cluster拓扑及状态：)"
    failover_pattern = r"\ * Coor Cluster Failover信息：\n(.*?)(?=\ * GBase版本号：)"
    coor_text = extract_blocks(text, coor_pattern)
    cron_line = extract_blocks(coor_text, cron_pattern)
    cron_line_new = re.sub(r'\n\s*\*$', '', cron_line.strip())
    ma_one = extract_coordinator1_ip(coor_text)
    system_name = get_cluster_name(text)
    failover_line = extract_blocks(coor_text, failover_pattern)
    failover_line_new = re.sub(r'\n\s*\*$', '', failover_line.strip())
    gbase_version = extract_gbase_version(coor_text)
    sys_data = [system_name, ma_one, gbase_version, cron_line_new, failover_line_new]

    cursor.executemany("INSERT OR IGNORE INTO sys_clusters (system_name, ma_one_ip, gbase_version, crontab_always, failover_info) VALUES (?, ?, ?, ?, ?)", [sys_data])

    conn.commit()
    conn.close()
    print(f"[✓] 成功写入集群信息到数据库")


# === 从文件中提取信息 ===
def operating_file(file_path):
    """对文件进行操作的示例函数"""
    # 这里可以添加对文件内容的处理逻辑
    print(f"[✓] 正在处理文件: {file_path}")
    file_content = read_file(file_path)

    # 提取特定块内容
    insert_ip_to_db(file_content, DB_PATH)
    get_machine_info(file_content, DB_PATH)
    get_machine_using(file_content, DB_PATH)
    get_cluster_disk_using(file_content, DB_PATH)
    get_cluster_process(file_content, DB_PATH)
    get_cluster_logs(file_content, DB_PATH)
    get_auto_start(file_content, DB_PATH)
    get_cluster_variables(file_content, DB_PATH)
    get_data_cluster_using(file_content, DB_PATH)
    get_instances(file_content, DB_PATH)
    get_sys_cluster(file_content, DB_PATH)


# === 根据每个系统名，进行巡检处理 开始 ===
def inspection_mppsystem(db_path):
    """从 sys_clusters 表中读取每个系统名"""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    cursor.execute("SELECT system_name FROM sys_clusters")
    rows = cursor.fetchall()

    for row in rows:
        system_name = row[0]
        print(f"正在处理：{system_name}")
        each_auto_inspection(system_name,DB_PATH)

    conn.close()


# === 总体运行情况 === 
def operational_status(system_name,db_path):
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    cursor.execute("select count(*) from machines where system_name = ?", (system_name,))
    row = cursor.fetchone()
    if row:
        all_nodes_count = row[0]

    cursor.execute("select count(*) from machines where system_name = ? and cluster_name = 'coor'", (system_name,))
    row = cursor.fetchone()
    if row:
        coor_nodes_count = row[0]

    data_nodes_count = all_nodes_count - coor_nodes_count

    cursor.execute("select gbase_version from sys_clusters where system_name = ?", (system_name,))
    row = cursor.fetchone()
    if row:
        gbase_version = row[0]    

    cursor.execute("select cluster_state,cluster_mode from data_clusters where system_name = ? limit 1", (system_name,))
    row = cursor.fetchone()
    if row:
        cluster_state, cluster_mode = row  # 拆包两个值
    else:
        print("没有找到记录")

    cursor.execute("""
        WITH counts AS (
            SELECT os_version , COUNT(*) AS count
            FROM machines where system_name = ?
            GROUP BY os_version 
        )
        SELECT os_version
        FROM counts
        WHERE count = (SELECT MAX(count) FROM counts)
                    """,(system_name,))
    row = cursor.fetchone()
    if row:
        platform = row[0]
    else:
        print("没有找到记录")

    cursor.execute("""
        WITH counts AS (
            SELECT cpu_model_name, cpu_logic_core, cpu_physical_core , COUNT(*) AS count
            FROM machines where system_name = ?
            GROUP BY cpu_logic_core, cpu_physical_core
        )
        SELECT cpu_model_name, cpu_logic_core, cpu_physical_core
        FROM counts
        WHERE count = (SELECT MAX(count) FROM counts)
                    """,(system_name,))
    row = cursor.fetchone()
    if row:
        cpu_model_name, cpu_logic_core, cpu_physical_core = row
    else:
        print("没有找到记录")
    match = re.search(r'model name\s*:\s*(.+)', cpu_model_name)
    if match:
        model_name = match.group(1)
        # print("提取到的 model name:", model_name)
    else:
        print("没有匹配到")

    cursor.execute("""
        WITH counts AS (
            SELECT mem_total, swap_total, COUNT(*) AS count
            FROM machine_using where system_name = ? and cluster_name = 'coor'
            GROUP BY mem_total, swap_total
        )
        SELECT mem_total, swap_total
        FROM counts
        WHERE count = (SELECT MAX(count) FROM counts)
                    """,(system_name,))
    row = cursor.fetchone()
    if row:
        cmem, cswap = row
    else:
        print("没有找到记录")

    cursor.execute("""
        WITH counts AS (
            SELECT mem_total, swap_total, COUNT(*) AS count
            FROM machine_using where system_name = ? and cluster_name <> 'coor'
            GROUP BY mem_total, swap_total
        )
        SELECT mem_total, swap_total
        FROM counts
        WHERE count = (SELECT MAX(count) FROM counts)
                    """,(system_name,))
    row = cursor.fetchone()
    if row:
        nmem, nswap = row
    else:
        print("没有找到记录")

    cursor.execute("""
select sum(disk_total), sum(disk_used ) from clusters_disk_using where system_name = ? and cluster_name <> 'coor'
                    """,(system_name,))
    row = cursor.fetchone()
    if row:
        ndisk_total, ndisk_used = row
    else:
        print("没有找到记录")
    ndisk_per = (ndisk_used / ndisk_total) * 100
    percent_4 = round(ndisk_per, 4)
    ndisk_used_tb = ndisk_used / 1024 / 1024 / 1024
    ndisk_used_tb_4 = round(ndisk_used_tb, 4)

    cursor.execute("""
SELECT sum(databases_count), sum(tables_count), sum(views_count), sum(procs_count), sum(funcs_count) from data_clusters dc where system_name = ?
                    """,(system_name,))
    row = cursor.fetchone()
    if row:
        ndbc,ntbc,nviewc,nprocc,nfuncc = row
    else:
        print("没有找到记录")

    conn.close()
    print("=============== 总体运行情况 ==================")
    print(f"""
Name:\t{system_name}\nNodes:\t{all_nodes_count},{coor_nodes_count},{data_nodes_count}\nRelese:\t{gbase_version}
State:\t{cluster_state}\nMode:\t{cluster_mode}\nPlatform:\t{platform}\nCPU:\tModel Name: {model_name}, Count: {cpu_physical_core}, Cores: {cpu_logic_core}
CoorMemTotal:\tMEM: {cmem}, Swap: {cswap}\nNodeMemTotal:\tMEM: {nmem}, Swap: {nswap}
DBSize:\t{ndisk_used_tb_4} ({percent_4}%)
DBUsed:\tDB num: {ndbc}, Tab num: {ntbc}, View num: {nviewc}, PROC num: {nprocc}, FUNC num: {nfuncc}
        """)
    report_data['ALL_NODE_C'] = f"ALL: {all_nodes_count}\nGC: {coor_nodes_count}\nGN: {data_nodes_count}"
    report_data['RELESE'] = gbase_version
    report_data['GSTATE'] = cluster_state
    report_data['GMODE'] = cluster_mode
    report_data['PLATFORM'] = platform
    report_data['CPU'] = f"{model_name}\n{cpu_physical_core}\n{cpu_logic_core}"
    report_data['MEMORY'] = f"Coor: {cmem}, Data: {nmem}\nCoor: {cswap}, Data: {nswap}"
    report_data['DBSIZE'] = f"{ndisk_used_tb_4}\n({percent_4}%)"
    report_data['DNUM'] = ndbc
    report_data['TNUM'] = ntbc
    report_data['VNUM'] = nviewc
    report_data['PNUM'] = nprocc
    report_data['FNUM'] = nfuncc



# === 空间可用性 ===
def data_cluster_used(system_name,db_path):
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("""
select sum(disk_total), sum(disk_used ) from clusters_disk_using where system_name = ? and cluster_name <> 'coor'
                    """,(system_name,))
    row = cursor.fetchone()
    if row:
        ndisk_total, ndisk_used = row
    else:
        print("没有找到记录")
    ndisk_per = (ndisk_used / ndisk_total) * 100
    percent_4 = round(ndisk_per, 4)
    ndisk_total_tb = ndisk_used / 1024 / 1024 / 1024
    ndisk_total_tb_4 = round(ndisk_total_tb, 4)
    ndisk_a = ndisk_total_tb * 0.8
    ndisk_a_4 = round(ndisk_a, 4)
    ndisk_used_tb = ndisk_used / 1024 / 1024 / 1024
    ndisk_used_tb_4 = round(ndisk_used_tb, 4)

    cursor.execute("""
select count(*) from machines where system_name = ? and cluster_name <> 'coor';
                    """,(system_name,))
    row = cursor.fetchone()
    if row:
        node_count = row[0]
    else:
        print("没有找到记录")


    conn.close()
    print("=============== 集群空间可用性 ==================")
    if ndisk_per > 80:
        print(f"⚠️ {system_name}: 集群空间大于80%，建议清理空间或者扩容")
        report_data["ALARM_DISK_USEING"] = "集群空间大于80%，建议清理空间或者扩容"
    else:
        report_data["ALARM_DISK_USEING"] = "集群空间小于80%，集群空间使用正常"

    print(f"""
集群共有{node_count}个数据节点，合计{ndisk_total_tb_4}TB存储空间，GBase集群实际可存储空间约为{ndisk_a_4}TB(GBase有效存储空间=总空间*80% )，目前已使用约{ndisk_used_tb_4}TB，约占总空间的{percent_4}%。
        """)
    dft = get_disk_from_db(system_name,db_path)
    print(dft)


    report_data['NODE_COUNT'] = node_count
    report_data['NDOSK_T'] = ndisk_total_tb_4
    report_data['NDOSK_A'] = ndisk_a_4
    report_data['NDISK_U'] = ndisk_used_tb_4
    report_data['PERCENT'] = percent_4
    # report_data['FNUM'] = nfuncc
    return dft


def write_df_to_table(doc: _Document, df: pd.DataFrame, table_index: int):
    """
    将 pandas.DataFrame 写入 Word 文档中的指定表格。
    保留表头，使用模板的第二行作为格式样式，将其删除后添加数据行。
    所有单元格内容统一设为小六号字（约 5pt）。
    
    :param doc: 已打开的 docx.Document 对象
    :param df: pandas.DataFrame，表格数据
    :param table_index: 要写入的表格在文档中的索引（从 0 开始）
    """
    tables = doc.tables
    if table_index >= len(tables):
        raise IndexError(f"文档中不存在索引为 {table_index} 的表格")

    table: Table = tables[table_index]

    if len(table.rows) < 2:
        raise ValueError("目标表格必须包含至少两行：表头和一个空模板行")

    # 复制第二行作为样式模板，然后删除它
    template_row = copy.deepcopy(table.rows[1]._tr)
    table._tbl.remove(table.rows[1]._tr)

    num_cols = len(table.rows[0].cells)

    # 删除旧数据行
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    # 添加新数据
    for _, row in df.iterrows():
        new_tr = copy.deepcopy(template_row)
        table._tbl.append(new_tr)
        new_row_cells = table.rows[-1].cells

        for i in range(min(num_cols, len(row))):
            cell = new_row_cells[i]
            cell.text = ""  # 清空原模板内容

            # 插入带样式的段落
            para = cell.paragraphs[0]
            run = para.add_run(str(row.iloc[i]))
            run.font.size = Pt(8)  # 字体大小 8 

# === 获取机器空间 ===
def get_disk_from_db(system_name,db_path):
    conn = sqlite3.connect(db_path)
    # cursor = conn.cursor()
#     cursor.execute("""
# select ROW_NUMBER() OVER () AS row_num, m.hostname, m.ip_address, mu.disk_size, mu.disk_used, mu.disk_avail, mu.disk_use_per
# from machines m join machine_using mu on m.ip_address = mu.ip_address
# where m.system_name = ? and m.cluster_name <> 'coor';
#                     """,(system_name,))
    sql = """
select ROW_NUMBER() OVER () AS row_num, m.hostname, m.ip_address, mu.disk_size, mu.disk_used, mu.disk_avail, mu.disk_use_per
from machines m join machine_using mu on m.ip_address = mu.ip_address
where m.system_name = ? and m.cluster_name <> 'coor';
                    """

    df = pd.read_sql_query(sql, conn, params=(system_name,))

    conn.close()
    return df


def extract_hostname_number(host):
    match = re.search(r'(\d+)$', host)
    return int(match.group(1)) if match else float('inf')

def check_row(row):
    if row["cluster_name"] == "coor":
        required = ["gclusterd", "gcrecover", "gcmonit", "gcmmonit"]
    else:
        required = ["gcmonit", "gcmmonit", "gbased", "gc_sync_server"]
    
    missing = [col for col in required if row[col] != 1]
    if missing:
        return f"{row['hostname']} ({row['ip']}): 缺少组件 {', '.join(missing)}"
    return None

def cluster_process_from_db(system_name,db_path):
    conn = sqlite3.connect(db_path)
    # === Step 1: 从数据库中读取数据 ===
    sql = """
    select m.hostname, cp.cluster_name, cp.ip_address, cp.process_cmd from clusters_process cp 
    join machines m on m.ip_address = cp.ip_address where cp.system_name = ?
    """
    df = pd.read_sql_query(sql, conn, params=(system_name,))
    conn.close()

    # === Step 2: 提取组件名（路径最后一级命令名） ===
    df["component"] = df["process_cmd"].str.extract(r'/([^/\s]+)(?:\s|$)')

    # === Step 3: 去重（同一个IP的同一组件只统计一次） ===
    df_unique = df[["hostname","cluster_name", "ip_address", "component"]].drop_duplicates()

    # === Step 4: 生成透视表（每个IP一行，每个组件一列，出现为1，未出现为0） ===
    df_unique["value"] = 1
    result = df_unique.pivot_table(index=["hostname", "cluster_name", "ip_address"], 
                                    columns="component", 
                                    values="value", 
                                    aggfunc="max", 
                                    fill_value=0)

    # === Step 5: 可选：重排列顺序，或填补缺失组件 ===
    all_components = ["gclusterd", "gcrecover", "gcmonit", "gcmmonit", "gbased", "gc_sync_server"]
    for comp in all_components:
        if comp not in result.columns:
            result[comp] = 0  # 如果某组件缺失就补0

    # 重排列顺序
    # print(result)
    result = result[all_components]

    # === Step 6: 重置索引、显示或导出 ===
    result = result.reset_index()
    # result['hostname_number'] = result['hostname'].apply(extract_hostname_number)
    # result = result.sort_values(by=['hostname_number'])
    # result = result.drop(columns='hostname_number')  # 排序后可删除辅助列

    print("=============== 集群进程诊断 ==================")
    
    result["alarm"] = result.apply(check_row, axis=1)
    alarms = result[result["alarm"].notnull()][["hostname", "ip_address", "alarm"]]
    if not alarms.empty:
        print(f"以下主机存在缺失组件：")
        print(alarms.to_string(index=False))
        report_data["ALARM_PROCE"] = "主机存在缺失组件"
    else:
        print(f"所有主机组件均部署完整。")
        report_data["ALARM_PROCE"] = "所有主机组件均部署完整"

    print(result)
    return result

def extract_log_type(path):
    if "system" in path.lower():
        return "system"
    elif "express" in path.lower():
        return "express"
    elif "gcrecover" in path.lower():
        return "gcrecover"
    elif "gc_sync_server" in path.lower():
        return "gc_sync_server"
    elif "core" in path.lower():
        return "core"
    elif "dump" in path.lower():
        return "dump"
    elif "loader_logs" in path.lower() or "loader" in path.lower():
        return "loader_logs"
    else:
        return None

def size_to_gb(size_str):
    """将 '5.3M'、'731M' 等字符串转换为 GB 数值"""
    try:
        if size_str == "0" or pd.isna(size_str):
            return 0
        match = re.match(r"([\d\.]+)([KMGTP])", size_str.upper())
        if not match:
            return 0
        num, unit = float(match.group(1)), match.group(2)
        factor = {"K": 1/1024/1024, "M": 1/1024, "G": 1, "T": 1024, "P": 1024*1024}
        return num * factor.get(unit, 0)
    except:
        return 0

def check_log_size_alerts(df, log_columns=None, threshold=800, verbose=True):
    """
    检查 DataFrame 中指定日志列是否超过 threshold（GB），并输出告警。

    参数:
        df           : pandas DataFrame（必须包含 hostname 和 ip_address）
        log_columns  : 要检查的列名列表（默认检查 7 个标准列）
        threshold    : GB 阈值，默认800
        verbose      : 是否打印告警信息，默认True

    返回:
        原始DataFrame的副本，附带 *_gb 列 和 log_alarm 列
    """
    if log_columns is None:
        log_columns = ["system", "express", "gcrecover", "gc_sync_server", "dump", "core", "loader_logs"]
    
    df = df.copy()  # 不修改原始表

    # 转换每列为 GB，并创建 *_gb 列
    for col in log_columns:
        gb_col = col + "_gb"
        df[gb_col] = df[col].apply(size_to_gb)

    # 生成告警信息列
    def get_oversized(row):
        return ", ".join([col for col in log_columns if row[col + "_gb"] > threshold])

    df["log_alarm"] = df.apply(get_oversized, axis=1)

    # 输出告警
    if verbose:
        alerts = df[df["log_alarm"] != ""]
        if not alerts.empty:
            print("⚠️ 以下主机存在日志文件超过 {} GB：\n".format(threshold))
            print(alerts[["hostname", "ip_address", "log_alarm"]].to_string(index=False))
            report_data["ALARM_LOGS_SIZE"] = f"存在日志文件超过 {threshold} GB"
        else:
            print("✅ 所有日志文件都未超过 {} GB".format(threshold))
            report_data["ALARM_LOGS_SIZE"] = f"日志文件正常"

    return df


def cluster_logs_from_db(system_name,db_path):
    conn = sqlite3.connect(db_path)
    # === Step 1: 从数据库中读取数据 ===
    sql = """
    select cl.cluster_name, m.hostname, cl.ip_address, cl.log_used, cl.log_path from clusters_logs cl 
    join machines m on m.ip_address = cl.ip_address where cl.system_name = ?
    """
    df = pd.read_sql_query(sql, conn, params=(system_name,))
    conn.close()
    df["log_type"] = df["log_path"].apply(extract_log_type)
    # === Step 2: 清理数据（去掉未知类型）===
    df = df.dropna(subset=["log_type"])

    # === Step 3: 透视表格（pivot） ===
    pivot = df.pivot_table(index=["hostname", "ip_address"], 
                       columns="log_type", 
                       values="log_used", 
                       aggfunc="first", 
                       fill_value="0")
    # === Step 4: 确保所有日志列都在（即使为空） ===
    all_logs = ["system", "express", "gcrecover", "gc_sync_server", "dump", "core", "loader_logs"]
    for log in all_logs:
        if log not in pivot.columns:
            pivot[log] = "0"

    # === Step 5: 重排序并重置索引 ===
    pivot = pivot[all_logs].reset_index()
    # === Step 6: 判断大于800G的 ===
    check_log_size_alerts(pivot, threshold=800)
    # === Step 7: 输出结果 ===
    print("=============== 集群日志清理诊断 ==================")
    print(pivot)
    return pivot



def check_component_state(df, component_cols=None, datastate_col="datastate", verbose=True):
    """
    检查组件状态是否异常。
    
    - 组件字段（如 gcware）值为 'OPEN' 或 None 视为正常
    - DataState 为 '0' 或 None 视为正常

    参数：
        df             : 原始 DataFrame，必须包含 NodeName 和 IpAddress
        component_cols : 要检查是否为 OPEN 的列，默认：['gcware', 'gcluster', 'gnode', 'syncserver']
        datastate_col  : 要检查是否为 '0' 的列名，默认：'DataState'
        verbose        : 是否打印告警信息

    返回：
        带 component_alarm 列的 DataFrame
    """
    if component_cols is None:
        component_cols = ['gcware', 'gcluster', 'gnode', 'syncserver']
    
    df = df.copy()

    def get_alarm(row):
        problems = []
        for col in component_cols:
            val = row.get(col)
            if val is not None and str(val).strip().upper() != "OPEN":
                problems.append(col)
        # DataState 允许为 "0" 或 None
        val = row.get(datastate_col)
        if val is not None and str(val).strip() != "0":
            problems.append(datastate_col.lower())  # 输出字段统一小写
        return ", ".join(problems) if problems else ""

    df["component_alarm"] = df.apply(get_alarm, axis=1)

    if verbose:
        alerts = df[df["component_alarm"] != ""]
        if not alerts.empty:
            print("⚠️ 以下节点存在异常组件状态：\n")
            print(alerts[["namenode", "ip_address", "component_alarm"]].to_string(index=False))
            report_data["ALARM_INSTANCE"] = "节点存在异常组件状态"
        else:
            print("✅ 所有组件状态正常")
            report_data["ALARM_INSTANCE"] = "所有组件状态正常"
    
    return df

def cluster_instance_from_db(system_name,db_path):
    conn = sqlite3.connect(db_path)
    # === Step 1: 从数据库中读取数据 ===
    sql = """
    select namenode, ip_address, gcware, gcluster, gnode, syncserver, datastate  from instances where system_name = ?
    """
    df = pd.read_sql_query(sql, conn, params=(system_name,))
    conn.close()
    check_component_state(df)
    # print(result) 
    print("=============== 集群实例诊断 ==================")
    print(df)
    return df


def process_service_logs(df):
    def extract_service_type(cmd):
        if "gcware_services" in cmd:
            return "gcware_services"
        elif "gcluster_services" in cmd:
            return "gcluster_services"
        return None

    df = df.copy()
    df["service"] = df["process_start"].apply(extract_service_type)
    df = df[df["service"].notna()].copy()
    df["value"] = 1

    pivot = df.pivot_table(
        index=["cluster_name", "hostname", "ip_address"],
        columns="service",
        values="value",
        aggfunc="max",
        fill_value=0
    ).reset_index()

    for col in ["gcware_services", "gcluster_services"]:
        if col not in pivot.columns:
            pivot[col] = 0

    def check_alarm(row):
        alarms = []
        if row["cluster_name"] == "coor":
            if row["gcware_services"] != 1:
                alarms.append("coor 节点缺少 gcware 服务")
            if row["gcluster_services"] != 1:
                alarms.append("coor 节点缺少 gcluster 服务")
        else:
            if row["gcluster_services"] != 1:
                alarms.append("非 coor 节点缺少 gcluster 服务")
            if row["gcware_services"] == 1:
                alarms.append("非 coor 节点不应启用 gcware 服务")
        return ", ".join(alarms)

    pivot["service_alarm"] = pivot.apply(check_alarm, axis=1)

    # 明确列顺序
    column_order = [
        "cluster_name", "hostname", "ip_address",
        "gcware_services", "gcluster_services",
        "service_alarm"
    ]
    for col in column_order:
        if col not in pivot.columns:
            pivot[col] = ""
    pivot = pivot[column_order]

    return pivot


def cluster_auto_start_from_db(system_name,db_path):
    conn = sqlite3.connect(db_path)
    # === Step 1: 从数据库中读取数据 ===
    sql = """
    select a.cluster_name, m.hostname, a.ip_address, a.process_start from auto_start a 
    join machines m on m.ip_address = a.ip_address where a.system_name = ?
    """
    df = pd.read_sql_query(sql, conn, params=(system_name,))
    conn.close()
    # final_df = process_service_logs(df)
    # print(final_df.to_string(index=False))
    result = process_service_logs(df)
    alerts = result[result["service_alarm"] != ""]
    if not alerts.empty:
        print("⚠️ 自启动服务异常，异常节点：")
        print(alerts[["hostname", "ip_address", "service_alarm"]].to_string(index=False))
        report_data["ALARM_AUTO_START"] = "自启动服务异常"
    else:
        print("✅ 所有服务自启动状态正常")
        report_data["ALARM_AUTO_START"] = "服务自启动状态正常"
    
    print("=============== 集群自启动诊断 ==================")
    print(result.to_string(index=False))
    return result


def smart_convert(value):
    """
    智能转换为数值（单位转字节），无法转换返回 None。
    """
    if pd.isna(value):
        return None
    try:
        str_val = str(value).strip()
        # 普通数字
        if re.match(r'^\d+(\.\d+)?$', str_val):
            return float(str_val)
        # 单位数字（如 30G, 1.2T）
        match = re.match(r'^([\d.]+)\s*([KMGTP])B?$', str_val, re.IGNORECASE)
        if match:
            num = float(match.group(1))
            unit = match.group(2).upper()
            unit_map = {
                'K': 1024,
                'M': 1024**2,
                'G': 1024**3,
                'T': 1024**4,
                'P': 1024**5
            }
            return num * unit_map[unit]
    except:
        pass
    return None

def check_cluster_params(df):
    if df.empty:
        print("⚠️ 没有可用的参数数据，请检查输入数据源。")
        return pd.DataFrame()

    df = df.copy()
    
    # 转换
    df["ref_val"] = df["var_reference"].apply(smart_convert)
    df["act_val"] = df["var_actual"].apply(smart_convert)

    # 告警判断
    def alarm_info(row):
        if pd.isna(row["ref_val"]):
            return "参考值为空或无效"
        if pd.isna(row["act_val"]):
            return "实际值为空或格式无法解析"
        if abs(row["ref_val"] - row["act_val"]) > 1e-6:
            return "实际值与参考值不一致"
        return ""

    df["告警说明"] = df.apply(alarm_info, axis=1)
    alerts = df[df["告警说明"] != ""]

    if alerts.empty:
        print("✅ 所有集群参数均正常，无异常告警。")
        report_data["ALARM_VARIABLES"] = "集群参数均正常"
    else:
        print("⚠️ 以下集群参数存在异常：")
        print(alerts[["ip_address", "var_name", "var_reference", "var_actual", "告警说明"]].to_string(index=False))
        report_data["ALARM_VARIABLES"] = "集群参数存在异常"
    
    return alerts

def cluster_variables_from_db(system_name,db_path):
    conn = sqlite3.connect(db_path)
    # === Step 1: 从数据库中读取数据 ===
    sql = """
    select ip_address, var_name, var_reference, var_actual from cluster_variables where system_name = ?
    """
    dfa = pd.read_sql_query(sql, conn, params=(system_name,))
    conn.close()
    print("=============== 集群参数诊断 ==================")
    check_cluster_params(dfa)

def extract_command_path(line):
    parts = line.split()
    for i, part in enumerate(parts):
        if part.startswith('/') or part.startswith('sh'):
            return ' '.join(parts[i:])  # 命令部分
    return ''

def check_crontab_entries(reference_text, actual_text, similarity_threshold=0.85):
    ref_lines = [line.strip() for line in reference_text.strip().splitlines() if line.strip()]
    actual_lines = [line.strip() for line in actual_text.strip().splitlines() if line.strip()]

    missing = []
    warnings = []

    for ref_line in ref_lines:
        if ref_line in actual_lines:
            continue

        ref_cmd = extract_command_path(ref_line)
        best_match = None
        highest_ratio = 0

        for actual_line in actual_lines:
            act_cmd = extract_command_path(actual_line)
            # 基于命令部分比较相似度
            ratio = difflib.SequenceMatcher(None, ref_cmd, act_cmd).ratio()
            if ratio > highest_ratio:
                highest_ratio = ratio
                best_match = actual_line

        if highest_ratio >= similarity_threshold:
            warnings.append((ref_line, best_match, highest_ratio))
        else:
            missing.append(ref_line)

    lines = []
    if missing:
        report_data["ALARM_CRON"] = "定时任务有缺失"
        lines.append("❌ 缺失的定时任务（完全未找到）：")
        for line in missing:
            lines.append(f"  MISSING: {line}")

    if warnings:
        report_data["ALARM_CRON"] = "定时任务有差异"
        lines.append("\n⚠️ 差异告警（存在相似项但不完全一致）：")
        for ref, match, score in warnings:
            lines.append(f"  WARNING: {ref}")
            lines.append(f"           ↳ 相似项: {match}")
            lines.append(f"           ↳ 相似度: {score:.2f}")

    if not missing and not warnings:
        lines.append("✅ 所有定时任务都完全匹配，无缺失无差异。")
        report_data["ALARM_CRON"] = "所有定时任务都完全正常"

    return "\n".join(lines)

def gcluster_script_from_db(system_name,db_path):
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("select gbase_version, crontab_always from sys_clusters where system_name = ?", (system_name,))
    row = cursor.fetchone()
    if row:
        gbase_version, crontab_always = row  # 拆包两个值
    else:
        print("没有找到记录")

    conn.close()
    pattern = r'\bsh\s+always\.sh\b'
    if re.search(pattern, crontab_always):
        print("✅ always.sh 脚本正在运行")
        report_data["ALARM_ALWAYS"] = "已启动always.sh运维脚本"
    else:
        print("⚠️ always.sh 脚本未运行")
        report_data["ALARM_ALWAYS"] = "未启动always.sh运维脚本"

    crontab_refence = """
30 1 * * * sh /opt/gbase_workspace/scripts/check_hole_lean/bin/run_test.sh
*/5 * * * * sh /opt/gbase_workspace/scripts/monitor/bin/monitor.sh
30 18 */15 * * cd /opt/gbase_workspace/scripts/monitor/logs;tar -czf abnormal.log.tar.gz abnormal.log
30 12 * * * sh /opt/gbase_workspace/scripts/delete_log/crontab_delete_logfile.sh
0 15 * * * sh /opt/gbase_workspace/scripts/inspection/inspection_gbase.sh
0 16 25 * * sh /opt/gbase_workspace/scripts/inspection/inspection_pro.sh
"""
    cron_text = check_crontab_entries(crontab_refence, crontab_always)
    report_data['CRON_TTEXT'] = cron_text

    crontab_text = []
    always_text = []    
    for line in crontab_always.strip().splitlines():
        if "always.sh" in line:
            always_text.append(line)
        else:
            crontab_text.append(line)

    # 拆分后的字符串结果
    crontab_result = "\n".join(crontab_text)
    always_result = "\n".join(always_text)
    report_data['CRON_RESULT'] = crontab_result
    report_data['ALWAYS_RESULT'] = always_result

    print(cron_text)
    print("=============== 集群定时任务和always脚本诊断 ==================")

def cluster_get_system_date(system_name,db_path):
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("select sc.ma_one_ip, f.filename  from sys_clusters sc join files f on sc.system_name = f.system_name where sc.system_name = ?", (system_name,))
    row = cursor.fetchone()
    if row:
        ma_one_ip, filename = row  # 拆包两个值
    else:
        print("没有找到记录")

    conn.close()

    match = re.search(r'\d{4}-\d{2}-\d{2}', filename)

    print("=============== 集群名、MA01和巡检时间获取 ==================")
    if match:
        date_str = match.group(0)
        year, month, _ = date_str.split("-")
        date_result = f"{year}年{int(month)}月"
    
    print(system_name,ma_one_ip,date_result)

    report_data['SYS_NAME'] = system_name
    report_data['MA_ONE_IP'] = ma_one_ip
    report_data['DATE_TEAR'] = date_result
    


def replace_placeholders(doc, replacements: dict):
    for paragraph in doc.paragraphs:
        for key, val in replacements.items():
            placeholder = f"{{{{{key}}}}}"
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(val))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, val in replacements.items():
                        placeholder = f"{{{{{key}}}}}"
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(val))




def each_auto_inspection(system_name,db_path):
    # 文本头
    cluster_get_system_date(system_name,db_path)
    # 总体运行情况
    operational_status(system_name,db_path)
    # 空间可用性
    df1 = data_cluster_used(system_name,db_path)

    # 检查进程
    df2 = cluster_process_from_db(system_name,db_path)
    cols_needed = ["ip_address", "gclusterd", "gcrecover", "gcmonit", "gcmmonit", "gbased", "gc_sync_server"]
    df_filtered = df2[cols_needed].copy()
    df_filtered.insert(0, "序号", range(1, len(df_filtered) + 1))
    df_filtered = df_filtered.replace(0, '')

    # 检查日志
    df3 = cluster_logs_from_db(system_name,db_path)
    all_logs = ["ip_address", "system", "express", "gcrecover", "gc_sync_server", "dump", "core", "loader_logs"]
    df_cluster_logs = df3[all_logs].copy()
    df_cluster_logs.insert(0, "序号", range(1, len(df_cluster_logs) + 1))

    # 检查实例
    df4 = cluster_instance_from_db(system_name,db_path)
    df4.insert(0, "序号", range(1, len(df4) + 1))
    df4 = df4.fillna('')

    # 检查自启动
    df5 = cluster_auto_start_from_db(system_name,db_path)
    c2 = ["ip_address", "gcware_services", "gcluster_services"]
    df_auto_start = df5[c2].copy()
    df_auto_start.insert(0, "序号", range(1, len(df_auto_start) + 1))
    df_auto_start = df_auto_start.replace(0, '')

    cluster_variables_from_db(system_name,db_path)

    gcluster_script_from_db(system_name,db_path)

    doc = Document(TEMPLATE_FILE)
    replace_placeholders(doc, report_data)
    # TEMPLATE_FILE_OUT
    ma01 = report_data["MA_ONE_IP"]
    inspection_date = report_data["DATE_TEAR"]
    # out_file_name = f"ZH-GBase8a集群-{system_name}系统-[{ma01}]-月度巡检报告-{inspection_date}.docx"
    out_file_name = f"ZH-GBase8a集群-{system_name}系统-灾备环境[{ma01}]-月度巡检报告-{inspection_date}.docx"
    file_path = os.path.join(TEMPLATE_FILE_OUT , out_file_name)

    write_df_to_table(doc, df1, table_index=2)
    write_df_to_table(doc, df_filtered, table_index=3)
    write_df_to_table(doc, df_cluster_logs, table_index=4)
    write_df_to_table(doc, df4, table_index=5)
    write_df_to_table(doc, df_auto_start, table_index=7)

    doc.save(file_path)
    print(f"✅ 报告已保存为：{file_path}")

# === 根据每个系统名，进行巡检处理 结束 ===
 
def main():
    initialize_project_directories()
    get_zip_file_path()
    print("== 自动化解压并写入数据库 ==")
    extract_zip(ZIP_FILE_PATH, EXTRACT_FOLDER)
    files = get_all_files(EXTRACT_FOLDER)
    init_database(DB_PATH)
    insert_files_to_db(DB_PATH, files, EXTRACT_FOLDER)
    print("== 现在开始处理每个文件内容 ==")
    process_each_file_from_db(DB_PATH)
    print("== 现在开始巡检每个系统内容 ==")
    inspection_mppsystem(DB_PATH)
    print("== 所有步骤执行完毕 ✅ ==")

    # # 显示所有列和行
    # pd.set_option('display.max_columns', None)  # 显示所有列
    # pd.set_option('display.max_rows', None)     # 显示所有行
    # pd.set_option('display.width', 0)           # 自动调整宽度以适应终端
    # pd.set_option('display.max_colwidth', None) # 显示所有列内容


if __name__ == '__main__':
    # main()
    with open("output.txt", "w", encoding="utf-8") as f:
        sys.stdout = f  # 重定向 print 到文件
        main()
        sys.stdout = sys.__stdout__  # 可选：恢复标准输出

    print("✅ 所有输出已写入 output.txt")